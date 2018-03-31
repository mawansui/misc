from docx import Document

def generate_table():
	# создаём объект документа
	document = Document()

	# откроем файл со всеми данными
	all_lines = []
	with open('final_data.txt') as file:
		all_lines = file.read().split("\n")[:-1] # чтобы без пустой линии

	# номер пробы и температура

	message = ("Слушай, извини, совсем забыл спросить – при какой температуре "
			   "проводился опыт? > ")
	temperature = float(input(message))

	# создаём объект таблицы
	table = document.add_table(rows = 1, cols = 7)

	# вытаскиваем объекты ячеек заголовка таблица и присваиваем им значения
	header_cells = table.rows[0].cells
	header_cells_titles = ["№ опыта", "T, ºC", "t, сек", "G, г", "V, мл", 
						   "Ge = G/b", "1/(1 – p)"]
	for cell, title in zip(header_cells, header_cells_titles):
		cell.text = title

	counter = 0

	for data_line in all_lines:
		counter += 1
		row_cells = table.add_row().cells
		row_cells[0].text = str(counter)
		row_cells[1].text = str(temperature)
		for x in range(2, 7):
			row_cells[x].text = data_line.split(" ")[x-2]

	table.style = "Table Grid"

	document.add_paragraph("\nГрафик зависимости 1/(1-p) от времени:\n")

	document.add_picture('time_vs_rate_of_reaction.png')

	document.save('labwork_table.docx')
	print("Вордовская таблица с названем {} успешно создана!".format('labwork_table.docx'))